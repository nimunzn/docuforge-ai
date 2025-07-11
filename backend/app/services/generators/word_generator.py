from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from typing import Dict, Any, List, Optional
import os
import base64


class WordGenerator:
    def __init__(self):
        self.styles_applied = False
    
    def create_document(self, document_data: Dict[str, Any], options: Dict[str, Any] = None) -> bytes:
        """Create a Word document from document data"""
        options = options or {}
        
        # Create new document
        doc = Document()
        
        # Apply custom styles
        self._apply_document_styles(doc)
        
        # Add document title
        title = document_data.get('title', 'Document')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.name = 'Arial'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(24)
        
        # Add metadata if available
        metadata = document_data.get('metadata', {})
        if metadata.get('created_at'):
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Created: {metadata['created_at'][:10]}")
            date_run.font.size = Pt(10)
            date_run.font.italic = True
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.space_after = Pt(12)
        
        # Add sections
        sections = document_data.get('sections', [])
        for i, section in enumerate(sections):
            self._add_section(doc, section, options)
            
            # Add page break between sections if requested
            if options.get('page_break_between_sections', False) and i < len(sections) - 1:
                doc.add_page_break()
        
        # Add footer with page numbers
        self._add_footer(doc)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _apply_document_styles(self, doc: Document):
        """Apply custom styles to the document"""
        if self.styles_applied:
            return
        
        styles = doc.styles
        
        # Custom heading style
        try:
            heading_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.base_style = styles['Heading 1']
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(16)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)
            
            heading_para = heading_style.paragraph_format
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)
            heading_para.keep_with_next = True
        except:
            pass  # Style might already exist
        
        # Custom body style
        try:
            body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = styles['Normal']
            body_font = body_style.font
            body_font.name = 'Arial'
            body_font.size = Pt(11)
            
            body_para = body_style.paragraph_format
            body_para.space_after = Pt(6)
            body_para.line_spacing = 1.15
        except:
            pass  # Style might already exist
        
        self.styles_applied = True
    
    def _add_section(self, doc: Document, section: Dict[str, Any], options: Dict[str, Any]):
        """Add a section to the document"""
        section_title = section.get('title', '')
        section_content = section.get('content', '')
        
        # Add section heading
        if section_title:
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(section_title)
            heading_run.font.size = Pt(16)
            heading_run.font.bold = True
            heading_run.font.name = 'Arial'
            heading_run.font.color.rgb = RGBColor(0, 51, 102)
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)
        
        # Add section content
        if section_content:
            # Split content into paragraphs
            paragraphs = section_content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    
                    # Handle basic formatting
                    self._add_formatted_text(para, para_text.strip())
                    
                    # Apply body style
                    para.style = doc.styles['Normal']
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.line_spacing = 1.15
        
        # Add subsections if any
        subsections = section.get('subsections', [])
        for subsection in subsections:
            self._add_subsection(doc, subsection, options)
    
    def _add_subsection(self, doc: Document, subsection: Dict[str, Any], options: Dict[str, Any]):
        """Add a subsection to the document"""
        subsection_title = subsection.get('title', '')
        subsection_content = subsection.get('content', '')
        
        # Add subsection heading
        if subsection_title:
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(subsection_title)
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.name = 'Arial'
            heading_run.font.color.rgb = RGBColor(0, 51, 102)
            heading_para.space_before = Pt(10)
            heading_para.space_after = Pt(4)
        
        # Add subsection content
        if subsection_content:
            paragraphs = subsection_content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, para_text.strip())
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.left_indent = Inches(0.25)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to paragraph with basic markdown-like formatting"""
        # Simple formatting: **bold**, *italic*, `code`
        import re
        
        # Split by formatting patterns
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code text
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Normal text
                paragraph.add_run(part)
    
    def _add_footer(self, doc: Document):
        """Add footer with page numbers"""
        try:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Page "
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field
            run = footer_para.runs[0]
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
            
        except Exception as e:
            # If footer creation fails, continue without it
            pass
    
    def add_table(self, doc: Document, table_data: List[List[str]], headers: List[str] = None):
        """Add a table to the document"""
        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        
        if headers:
            rows += 1
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Add headers
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Add data
        start_row = 1 if headers else 0
        for i, row_data in enumerate(table_data):
            row_cells = table.rows[start_row + i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = str(cell_data)
        
        return table
    
    def add_image(self, doc: Document, image_path: str, width: Optional[float] = None, caption: str = None):
        """Add an image to the document"""
        try:
            if width:
                doc.add_picture(image_path, width=Inches(width))
            else:
                doc.add_picture(image_path)
            
            # Add caption if provided
            if caption:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Figure: {caption}")
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.space_after = Pt(12)
                
        except Exception as e:
            # If image addition fails, add placeholder text
            para = doc.add_paragraph()
            run = para.add_run(f"[Image: {caption or 'Image could not be loaded'}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    def create_template_document(self, template_type: str, title: str = None) -> bytes:
        """Create a template document based on type"""
        templates = {
            'proposal': {
                'title': title or 'Business Proposal',
                'sections': [
                    {'title': 'Executive Summary', 'content': '[Write a brief overview of your proposal]'},
                    {'title': 'Problem Statement', 'content': '[Describe the problem you are solving]'},
                    {'title': 'Proposed Solution', 'content': '[Explain your proposed solution]'},
                    {'title': 'Implementation Timeline', 'content': '[Outline the timeline for implementation]'},
                    {'title': 'Budget and Resources', 'content': '[Detail the budget and resources needed]'},
                    {'title': 'Expected Outcomes', 'content': '[Describe the expected results]'},
                    {'title': 'Conclusion', 'content': '[Summarize the key points and call to action]'}
                ]
            },
            'report': {
                'title': title or 'Report',
                'sections': [
                    {'title': 'Executive Summary', 'content': '[Provide a brief summary of the report]'},
                    {'title': 'Introduction', 'content': '[Introduce the topic and objectives]'},
                    {'title': 'Methodology', 'content': '[Explain the methods used]'},
                    {'title': 'Findings', 'content': '[Present the key findings]'},
                    {'title': 'Analysis', 'content': '[Analyze the findings]'},
                    {'title': 'Recommendations', 'content': '[Provide recommendations]'},
                    {'title': 'Conclusion', 'content': '[Summarize the conclusions]'}
                ]
            },
            'memo': {
                'title': title or 'Memorandum',
                'sections': [
                    {'title': 'TO:', 'content': '[Recipient names]'},
                    {'title': 'FROM:', 'content': '[Your name]'},
                    {'title': 'DATE:', 'content': '[Current date]'},
                    {'title': 'SUBJECT:', 'content': '[Subject line]'},
                    {'title': 'Purpose', 'content': '[State the purpose of the memo]'},
                    {'title': 'Background', 'content': '[Provide relevant background information]'},
                    {'title': 'Discussion', 'content': '[Present the main points]'},
                    {'title': 'Action Items', 'content': '[List specific actions required]'}
                ]
            }
        }
        
        template_data = templates.get(template_type, templates['proposal'])
        return self.create_document(template_data)